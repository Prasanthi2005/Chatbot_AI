# ==========================================
# exports/export_docx.py
# Export Chat History to DOCX
# ==========================================

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt


def export_chat_to_docx(username, chats, export_folder="exports"):
    """
    Export chat history to a DOCX file.

    Parameters:
        username (str): Username
        chats (list): List of chat rows/dictionaries
        export_folder (str): Folder to save DOCX

    Returns:
        str: Full path of generated DOCX file
    """

    os.makedirs(export_folder, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    filename = f"{username}_chat_{timestamp}.docx"

    filepath = os.path.join(export_folder, filename)

    document = Document()

    title = document.add_heading("AI Chatbot Chat History", level=1)
    title.style.font.size = Pt(20)

    document.add_paragraph(f"Username : {username}")
    document.add_paragraph(
        f"Generated : {datetime.now().strftime('%d-%m-%Y %I:%M:%S %p')}"
    )

    document.add_paragraph("-" * 60)

    if not chats:
        document.add_paragraph("No chat history available.")

    else:

        for index, chat in enumerate(chats, start=1):

            if isinstance(chat, dict):
                user_message = chat.get("user_message", "")
                bot_reply = chat.get("bot_reply", "")

            else:
                user_message = chat["user_message"]
                bot_reply = chat["bot_reply"]

            document.add_heading(f"Chat {index}", level=2)

            p1 = document.add_paragraph()
            p1.add_run("You:\n").bold = True
            p1.add_run(user_message)

            p2 = document.add_paragraph()
            p2.add_run("AI:\n").bold = True
            p2.add_run(bot_reply)

            document.add_paragraph("-" * 50)

    document.save(filepath)

    return filepath